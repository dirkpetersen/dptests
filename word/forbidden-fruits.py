#!/usr/bin/env python3

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.text.run import Run
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement
from docx.enum.dml import MSO_COLOR_TYPE # For copying font color accurately
# from docx.shared import Pt, RGBColor # Not strictly needed for copying but good to be aware of

# --- Helper function to copy font properties ---
def _copy_font_properties(source_font, target_font):
    """Copies font properties from source_font to target_font."""
    target_font.name = source_font.name
    target_font.size = source_font.size
    target_font.bold = source_font.bold
    target_font.italic = source_font.italic
    target_font.underline = source_font.underline
    target_font.strike = source_font.strike
    target_font.subscript = source_font.subscript
    target_font.superscript = source_font.superscript
    target_font.all_caps = source_font.all_caps
    target_font.small_caps = source_font.small_caps
    # For complex properties like color, check type before assigning
    if source_font.color and hasattr(source_font.color, 'type'):
        if source_font.color.type == MSO_COLOR_TYPE.RGB:
            target_font.color.rgb = source_font.color.rgb
        elif source_font.color.type == MSO_COLOR_TYPE.THEME:
            # This might require more complex handling if theme colors need to be resolved
            # For simplicity, we'll try to assign theme_color index.
            # If direct assignment isn't perfect, it's a limitation of simple copy.
            target_font.color.theme_color = source_font.color.theme_color
    # Add other properties as needed, e.g., highlight_color (though we set this separately)

# --- Core highlighting logic for a paragraph ---
def highlight_paragraph(paragraph, compiled_patterns):
    """
    Highlights terms in a paragraph based on compiled_patterns.
    compiled_patterns is a list of (regex_pattern, original_term_string) tuples,
    sorted by term length descending.
    """
    for pattern, _ in compiled_patterns: # original_term_string can be used for logging
        run_idx = 0
        while run_idx < len(paragraph.runs):
            run = paragraph.runs[run_idx]

            # Skip if run is already highlighted by a previous (longer/more specific) term
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                run_idx += 1
                continue

            if not run.text: # Skip empty runs
                run_idx += 1
                continue

            # re.split with a capturing group will include the matched delimiters in the result
            # e.g., pattern=(r'(\bterm\b)'), text='before term after' -> ['before ', 'term', ' after']
            try:
                # pattern already has re.IGNORECASE flag from compilation
                text_parts = pattern.split(run.text)
            except re.error: # Should not happen if patterns are pre-compiled correctly
                run_idx +=1
                continue


            if len(text_parts) == 1: # No match for this pattern in this run's text
                run_idx += 1
                continue

            # Match found. The run needs to be split.
            # text_parts = [before_match, match1, between_match1_match2, match2, after_match2, ...]

            # Store original run's properties to apply to new runs
            # We need to get these before modifying run.text, as it can affect style inheritance
            original_run_element = run._r # The lxml element of the run
            original_style = run.style.name if run.style else None

            # Create a temporary Run object to easily access font properties if run.text becomes empty
            # This is a bit of a workaround to ensure font properties are captured correctly
            # before run.text is modified, especially if parts[0] is empty.
            temp_doc = Document()
            temp_run = temp_doc.add_paragraph().add_run()
            _copy_font_properties(run.font, temp_run.font)


            # The first part (text before the first match) stays in the current run
            run.text = text_parts[0]
            # Re-apply original font properties and style, as setting run.text might reset them
            # or if text_parts[0] is empty, the run might lose its formatting context.
            _copy_font_properties(temp_run.font, run.font)
            if original_style:
                run.style = original_style

            num_inserted_runs_for_this_pattern_in_this_run = 0
            current_insert_anchor_r = original_run_element

            for i in range(1, len(text_parts)):
                part_text = text_parts[i]
                if not part_text: # Skip empty parts that re.split might produce
                    continue

                # Create new <w:r> element and a new Run object
                new_r_element = OxmlElement('w:r')
                current_insert_anchor_r.addnext(new_r_element) # Insert XML element after the current anchor
                new_run = Run(new_r_element, paragraph) # Create Run object linked to paragraph

                new_run.text = part_text
                _copy_font_properties(temp_run.font, new_run.font) # Copy formatting from original
                if original_style:
                    new_run.style = original_style

                is_match_part = (i % 2 == 1) # Matched parts are at odd indices in text_parts
                if is_match_part:
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                current_insert_anchor_r = new_r_element # Next insertion will be after this newly added run
                num_inserted_runs_for_this_pattern_in_this_run += 1

            # Advance run_idx by the number of runs effectively processed/created
            # (1 for the original run which is now the first part, + num_inserted_runs)
            run_idx += (1 + num_inserted_runs_for_this_pattern_in_this_run)
            # Deleting temp_doc and temp_run is not strictly necessary due to Python's GC,
            # but good practice if they held significant resources. Here, they are small.
            del temp_doc
            del temp_run


# --- Functions to load and process terms ---
def load_terms_from_file(filepath):
    """Loads terms from a file, one term per line."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            terms = [line.strip() for line in f if line.strip()]
        return terms
    except FileNotFoundError:
        print(f"Error: Terms file not found at {filepath}")
        return []

def generate_term_variations(term):
    """Generates variations for a term (original, hyphenated, no-space)."""
    variations = {term} # Use a set to keep unique variations
    if ' ' in term:
        variations.add(term.replace(' ', '-'))
        variations.add(term.replace(' ', ''))
    return list(variations)

def get_compiled_search_patterns(raw_terms):
    """
    Generates all variations, sorts them by length (desc),
    and compiles them into regex patterns for splitting.
    """
    all_variations_with_original = []
    for term in raw_terms:
        variations = generate_term_variations(term)
        for var in variations:
            all_variations_with_original.append(var) # Keep original term for context/logging if needed

    # Sort unique variations by length, longest first, to handle overlapping terms correctly.
    # E.g., "apple pie" should be processed before "apple".
    unique_variations = sorted(list(set(all_variations_with_original)), key=len, reverse=True)

    compiled_patterns = []
    for var_text in unique_variations:
        # Regex for whole-word match, case-insensitive. Capturing group for re.split().
        # Using lookarounds for word boundaries to be more robust with hyphens/special chars
        # (?<!\w) ensures not preceded by a word character.
        # (?!\w) ensures not followed by a word character.
        # This is often more reliable than \b with complex terms.
        # However, \b is simpler and usually sufficient. Let's stick to \b as per common practice.
        # If \b causes issues with hyphenated terms being treated as separate words,
        # then (?<!\w) and (?!\w) would be the alternative.
        # For "term-name", \bterm-name\b works.
        # For "term name", \bterm name\b works.
        pattern_str = r'(\b' + re.escape(var_text) + r'\b)'
        try:
            # Ensure case insensitive search is applied
            compiled_patterns.append((re.compile(pattern_str, re.IGNORECASE), var_text))
        except re.error as e:
            print(f"Warning: Could not compile regex for term '{var_text}': {e}")
    return compiled_patterns

# --- Main processing function for a document ---
def process_document(doc_path, compiled_patterns):
    """Opens a document, highlights terms, and saves it."""
    print(f"Processing document: {doc_path}...")
    try:
        doc = Document(doc_path)

        # Process paragraphs in the main body
        for para in doc.paragraphs:
            highlight_paragraph(para, compiled_patterns)

        # Process paragraphs in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        highlight_paragraph(para, compiled_patterns)

        # Note: Headers, footers, text boxes, etc., are not processed by default with doc.paragraphs.
        # Handling them would require accessing those parts of the document specifically.
        # For now, sticking to main body and tables as per typical interpretation.

        doc.save(doc_path)
        print(f"Finished processing and saved: {doc_path}")
    except Exception as e:
        print(f"Error processing document {doc_path}: {e}")

# --- Script entry point ---
def main():
    parser = argparse.ArgumentParser(description="Highlight terms in Word documents.")
    parser.add_argument("folder_path", type=str, help="Path to the folder containing .docx files.")
    parser.add_argument("terms_file", type=str, help="Path to the text file with terms to highlight (one per line).")

    args = parser.parse_args()

    folder = Path(args.folder_path)
    terms_filepath = Path(args.terms_file)

    if not folder.is_dir():
        print(f"Error: Folder not found at {args.folder_path}")
        return
    if not terms_filepath.is_file():
        print(f"Error: Terms file not found at {args.terms_file}")
        return

    raw_terms = load_terms_from_file(terms_filepath)
    if not raw_terms:
        print("No terms to highlight.")
        return

    compiled_patterns = get_compiled_search_patterns(raw_terms)
    if not compiled_patterns:
        print("No valid search patterns could be compiled.")
        return

    print(f"Looking for .docx files in: {folder.resolve()}")
    print(f"Loaded {len(raw_terms)} terms for case-insensitive highlighting")
    docx_files_found = False
    for doc_path in folder.glob("*.docx"):
        if doc_path.name.startswith("~"): # Skip temporary Word files
            continue
        docx_files_found = True
        process_document(doc_path, compiled_patterns)

    if not docx_files_found:
        print("No .docx files found in the specified folder.")

if __name__ == "__main__":
    main()


